"""Document and spreadsheet property + body-text extraction (untrusted input).

Three property parsers plus a **separately-bounded body-text pass**, each wrapped
with the same discipline as the ffprobe extractor: a file-size ceiling enforced
before parsing, no network access, and every field parsed defensively. On any
failure a DocumentError is raised carrying a message safe to store under
``_extract_error`` — the extract job never dies.

    * PDF  (pypdf)      — page count, title/author/subject/creator/producer,
                          encrypted flag, and (P3-T5) plain body text.
                          pypdf performs no network I/O. Encrypted PDFs that need
                          a password are reported as ``encrypted`` without a
                          decrypt attempt (never guess/brute-force) and yield NO
                          body text.
    * DOCX (python-docx)— core properties (title/author/created/modified/
                          subject/keywords) + paragraph count, and (P3-T5) the
                          concatenated paragraph text as body text.
    * XLSX (openpyxl)   — read_only + data_only: sheet names, sheet count, core
                          properties. Never loads cell data; NO body text.
    * TXT / MD          — no property parser; (P3-T5) a size-capped byte read
                          decoded errors="replace" as body text.

Zip-based formats (docx/xlsx) are a decompression-bomb vector. BEFORE either
library opens the archive, :func:`guard_decompression` inspects the zip central
directory (declared ``file_size`` vs ``compress_size``) and REJECTS a crafted
ratio-bomb — the total-uncompressed ceiling and the ratio gate are both checked
without decompressing a single member (cheap, and provably before any parse).

Body text is (a) capped at a character ceiling (``FILEARR_BODY_TEXT_MAX_CHARS``,
default 100_000 — snippets don't need novels), (b) whitespace-normalised, and
(c) stripped of C0/C1 control characters at STORE time (untrusted content — the
frontend never {@html}-renders it). ``body_text_truncated`` records whether the
cap (char OR page/read) clipped the content.

Emitted schema (all keys optional):
  PDF:   pages:int, title/author/subject/creator/producer:str,
         encrypted:bool, created/modified:str, body_text:str,
         body_text_truncated:bool
  DOCX:  title/author/subject:str, keywords:str, created/modified:str (ISO),
         paragraphs:int, revision:int, body_text:str, body_text_truncated:bool
  XLSX:  sheets:list[str], sheet_count:int, title/author/subject:str,
         created/modified:str (ISO)
  TXT/MD: body_text:str, body_text_truncated:bool
"""

from __future__ import annotations

import os
import re
import zipfile
from pathlib import PurePath
from typing import Any


class DocumentError(RuntimeError):
    """A document/spreadsheet could not be parsed. Message is safe to store."""


# ---------------------------------------------------------------------------
# Body-text + decompression-guard defaults (overridable via config at call time;
# the module constants keep the pure functions self-contained + unit-testable).
# ---------------------------------------------------------------------------
DEFAULT_BODY_TEXT_MAX_CHARS = 100_000
# Zip-bomb guard: reject a docx/xlsx whose members decompress past this TOTAL, or
# whose overall ratio exceeds RATIO once the payload is already non-trivial.
DEFAULT_DECOMPRESSED_MAX = 200 * 1024 * 1024  # 200 MiB total uncompressed
DEFAULT_DECOMPRESSION_RATIO = 100.0  # uncompressed:compressed
DEFAULT_DECOMPRESSION_RATIO_MIN_BYTES = 10 * 1024 * 1024  # 10 MiB

_WS_RE = re.compile(r"\s+")


def _clean_str(v: Any) -> str | None:
    if v is None:
        return None
    s = str(v).strip()
    return s or None


def _iso(v: Any) -> str | None:
    """datetime → ISO string; pass through non-empty strings; else None."""
    if v is None:
        return None
    iso = getattr(v, "isoformat", None)
    if callable(iso):
        try:
            return iso()
        except Exception:
            return None
    return _clean_str(v)


def _guard_size(path: str, max_bytes: int) -> None:
    try:
        size = os.path.getsize(path)
    except OSError as exc:
        raise DocumentError(f"cannot stat document: {exc}") from exc
    if size > max_bytes:
        raise DocumentError(f"document too large ({size} > {max_bytes} bytes)")


def guard_decompression(
    path: str,
    *,
    decompressed_max: int = DEFAULT_DECOMPRESSED_MAX,
    ratio_limit: float = DEFAULT_DECOMPRESSION_RATIO,
    ratio_min_bytes: int = DEFAULT_DECOMPRESSION_RATIO_MIN_BYTES,
) -> None:
    """Reject a decompression-bomb docx/xlsx BEFORE any parser opens it.

    Reads ONLY the zip central directory (no member is decompressed) and sums the
    declared ``file_size`` (uncompressed) and ``compress_size`` (on-disk). Rejects
    when the total uncompressed size exceeds ``decompressed_max`` OR the overall
    ratio exceeds ``ratio_limit`` once the payload is already larger than
    ``ratio_min_bytes`` (so an ordinary tiny, highly-compressible office file — a
    few KB that inflate 100x — is NOT falsely rejected; only a genuine bomb is).

    Raises DocumentError on rejection or on a corrupt/oversized central directory.
    This is the same size-ceiling-before-parse discipline P3-T13 will extend to
    archive members.
    """
    try:
        with zipfile.ZipFile(path) as zf:
            infos = zf.infolist()
    except zipfile.BadZipFile as exc:
        raise DocumentError(f"not a valid zip container: {exc}") from exc
    except OSError as exc:
        raise DocumentError(f"cannot open zip container: {exc}") from exc

    total_uncompressed = sum(int(i.file_size) for i in infos)
    total_compressed = sum(int(i.compress_size) for i in infos)

    if total_uncompressed > decompressed_max:
        raise DocumentError(
            "decompression guard: declared uncompressed size "
            f"{total_uncompressed} exceeds ceiling {decompressed_max} bytes"
        )
    if total_compressed > 0:
        ratio = total_uncompressed / total_compressed
        if ratio > ratio_limit and total_uncompressed > ratio_min_bytes:
            raise DocumentError(
                f"decompression guard: compression ratio {ratio:.1f}:1 exceeds "
                f"{ratio_limit:.0f}:1 at {total_uncompressed} uncompressed bytes"
            )


def _normalize_body_text(
    raw: str, max_chars: int, *, hard_stopped: bool = False
) -> tuple[str, bool]:
    """Sanitise + whitespace-normalise + char-cap untrusted body text.

    Drops C0/C1 control characters (whitespace controls collapse to a space),
    collapses whitespace runs to single spaces, strips ends, and truncates to
    ``max_chars``. Returns ``(clean_text, truncated)`` — ``truncated`` is True if
    the char cap clipped the text OR the caller already stopped short
    (``hard_stopped``, e.g. a page/read ceiling hit upstream). Mirrors
    ``errors.sanitize_error`` control-char discipline (injection defence).
    """
    out: list[str] = []
    for ch in raw:
        o = ord(ch)
        if o in (0x09, 0x0A, 0x0B, 0x0C, 0x0D):
            out.append(" ")
        elif o < 0x20 or 0x7F <= o <= 0x9F:
            continue  # strip other control chars (ANSI/NUL/etc.)
        else:
            out.append(ch)
    collapsed = _WS_RE.sub(" ", "".join(out)).strip()
    truncated = hard_stopped or len(collapsed) > max_chars
    if len(collapsed) > max_chars:
        collapsed = collapsed[:max_chars].rstrip()
    return collapsed, truncated


def extract_pdf(path: str, *, max_bytes: int) -> dict[str, Any]:
    _guard_size(path, max_bytes)
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    meta: dict[str, Any] = {}
    try:
        # strict=False: tolerate malformed PDFs instead of raising mid-parse.
        reader = PdfReader(path, strict=False)
        # is_encrypted is cheap and set before page access.
        encrypted = bool(getattr(reader, "is_encrypted", False))
        meta["encrypted"] = encrypted
        if encrypted:
            # Try the empty-password path many "encrypted" PDFs use; never guess
            # a real password. If it stays locked, report encrypted + stop before
            # any page/metadata access (which would raise FileNotDecrypted).
            unlocked = False
            try:
                result = reader.decrypt("")
                # pypdf returns a PasswordType enum: 0 == NOT_DECRYPTED.
                unlocked = bool(result) and int(getattr(result, "value", result)) != 0
            except Exception:
                unlocked = False
            if not unlocked:
                return meta
        try:
            meta["pages"] = len(reader.pages)
        except Exception:
            pass  # page tree unreadable — keep whatever else we have
        info = getattr(reader, "metadata", None)
        if info is not None:
            # Per-field defensive access (live class D). pypdf decodes each
            # DocumentInformation property lazily and a MALFORMED value throws on
            # access -- most infamously a non-conformant /CreationDate
            # ("2006/05/24 21:06") raising ValueError "Can not convert date". That
            # must NOT sink the whole document extract, so every field is read in
            # its own try; a failure just drops that one field.
            for out_key, attr in (
                ("title", "title"),
                ("author", "author"),
                ("subject", "subject"),
                ("creator", "creator"),
                ("producer", "producer"),
            ):
                try:
                    val = _clean_str(getattr(info, attr, None))
                except Exception:
                    val = None
                if val is not None:
                    meta[out_key] = val
            # Dates: prefer the typed (parsed) datetime, but when pypdf cannot
            # parse a malformed date, fall back to the RAW string rather than
            # discarding it (integrity > tidiness -- keep what the file states).
            for out_key, attr, raw_attr in (
                ("created", "creation_date", "creation_date_raw"),
                ("modified", "modification_date", "modification_date_raw"),
            ):
                val = None
                try:
                    val = _iso(getattr(info, attr, None))
                except Exception:
                    try:
                        val = _clean_str(getattr(info, raw_attr, None))
                    except Exception:
                        val = None
                if val is not None:
                    meta[out_key] = val
    except (PyPdfError, ValueError, OSError, EOFError) as exc:
        raise DocumentError(f"pypdf could not read PDF: {exc}") from exc
    except Exception as exc:  # any other pypdf internal error → safe message
        raise DocumentError(f"pypdf failed: {exc}") from exc
    return meta


def extract_docx(
    path: str,
    *,
    max_bytes: int,
    decompressed_max: int = DEFAULT_DECOMPRESSED_MAX,
    ratio_limit: float = DEFAULT_DECOMPRESSION_RATIO,
    ratio_min_bytes: int = DEFAULT_DECOMPRESSION_RATIO_MIN_BYTES,
) -> dict[str, Any]:
    _guard_size(path, max_bytes)
    # Decompression-bomb guard BEFORE python-docx opens the archive (P3-T5).
    guard_decompression(
        path,
        decompressed_max=decompressed_max,
        ratio_limit=ratio_limit,
        ratio_min_bytes=ratio_min_bytes,
    )
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    meta: dict[str, Any] = {}
    try:
        document = docx.Document(path)
    except PackageNotFoundError as exc:
        raise DocumentError(f"not a valid DOCX package: {exc}") from exc
    except Exception as exc:
        raise DocumentError(f"python-docx could not open file: {exc}") from exc

    try:
        cp = document.core_properties
        for out_key, attr in (
            ("title", "title"),
            ("author", "author"),
            ("subject", "subject"),
            ("keywords", "keywords"),
        ):
            val = _clean_str(getattr(cp, attr, None))
            if val is not None:
                meta[out_key] = val
        for out_key, attr in (("created", "created"), ("modified", "modified")):
            iso = _iso(getattr(cp, attr, None))
            if iso is not None:
                meta[out_key] = iso
        rev = getattr(cp, "revision", None)
        if isinstance(rev, int) and rev:
            meta["revision"] = rev
    except Exception:
        pass  # core properties are optional; keep going

    try:
        meta["paragraphs"] = len(document.paragraphs)
    except Exception:
        pass
    return meta


def extract_xlsx(
    path: str,
    *,
    max_bytes: int,
    decompressed_max: int = DEFAULT_DECOMPRESSED_MAX,
    ratio_limit: float = DEFAULT_DECOMPRESSION_RATIO,
    ratio_min_bytes: int = DEFAULT_DECOMPRESSION_RATIO_MIN_BYTES,
) -> dict[str, Any]:
    _guard_size(path, max_bytes)
    # Same decompression-bomb guard as docx (P3-T5 retrofit) — xlsx is a zip too.
    guard_decompression(
        path,
        decompressed_max=decompressed_max,
        ratio_limit=ratio_limit,
        ratio_min_bytes=ratio_min_bytes,
    )
    import openpyxl
    from openpyxl.utils.exceptions import InvalidFileException

    meta: dict[str, Any] = {}
    wb = None
    try:
        # read_only avoids loading the full worksheet; data_only returns cached
        # values (never evaluates formulas). keep_links=False: no external refs.
        wb = openpyxl.load_workbook(
            path, read_only=True, data_only=True, keep_links=False
        )
    except InvalidFileException as exc:
        raise DocumentError(f"not a valid XLSX file: {exc}") from exc
    except Exception as exc:
        raise DocumentError(f"openpyxl could not open workbook: {exc}") from exc

    try:
        names = list(wb.sheetnames)
        meta["sheets"] = names
        meta["sheet_count"] = len(names)
        props = getattr(wb, "properties", None)
        if props is not None:
            for out_key, attr in (
                ("title", "title"),
                ("author", "creator"),
                ("subject", "subject"),
            ):
                val = _clean_str(getattr(props, attr, None))
                if val is not None:
                    meta[out_key] = val
            for out_key, attr in (("created", "created"), ("modified", "modified")):
                iso = _iso(getattr(props, attr, None))
                if iso is not None:
                    meta[out_key] = iso
    except Exception as exc:
        raise DocumentError(f"openpyxl failed reading workbook structure: {exc}") from exc
    finally:
        if wb is not None:
            try:
                wb.close()
            except Exception:
                pass
    return meta


# ---------------------------------------------------------------------------
# Body-text pass (P3-T5) — independently bounded from property extraction.
# ---------------------------------------------------------------------------
def _pdf_body(path: str, *, max_chars: int, max_bytes: int) -> tuple[str, bool]:
    """Concatenated PDF page text, char-capped. Encrypted PDFs yield ("", False).

    Per-page try/except: one unreadable page loses only that page. Extraction
    stops as soon as the running length reaches ``max_chars`` (time/space box) —
    the char normalisation applies the final hard cap.
    """
    _guard_size(path, max_bytes)
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(path, strict=False)
        if bool(getattr(reader, "is_encrypted", False)):
            unlocked = False
            try:
                result = reader.decrypt("")
                unlocked = bool(result) and int(getattr(result, "value", result)) != 0
            except Exception:
                unlocked = False
            if not unlocked:
                return "", False  # never extract text from a locked PDF
        parts: list[str] = []
        total = 0
        stopped = False
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
            except Exception:
                continue  # one bad page loses only that page
            if text:
                parts.append(text)
                total += len(text)
            if total >= max_chars:
                stopped = True
                break
        return _normalize_body_text("\n".join(parts), max_chars, hard_stopped=stopped)
    except (PyPdfError, ValueError, OSError, EOFError) as exc:
        raise DocumentError(f"pypdf could not extract text: {exc}") from exc
    except Exception as exc:
        raise DocumentError(f"pypdf text extraction failed: {exc}") from exc


def _docx_body(path: str, *, max_chars: int, max_bytes: int) -> tuple[str, bool]:
    """Concatenated DOCX paragraph text, char-capped. Caller runs the guard first.

    Per-paragraph try/except; stops once the running length reaches ``max_chars``.
    """
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(path)
    except PackageNotFoundError as exc:
        raise DocumentError(f"not a valid DOCX package: {exc}") from exc
    except Exception as exc:
        raise DocumentError(f"python-docx could not open file: {exc}") from exc

    parts: list[str] = []
    total = 0
    stopped = False
    try:
        paragraphs = document.paragraphs
    except Exception as exc:
        raise DocumentError(f"python-docx could not read paragraphs: {exc}") from exc
    for para in paragraphs:
        try:
            text = para.text or ""
        except Exception:
            continue  # one bad paragraph loses only itself
        if text:
            parts.append(text)
            total += len(text)
        if total >= max_chars:
            stopped = True
            break
    return _normalize_body_text("\n".join(parts), max_chars, hard_stopped=stopped)


def _text_body(path: str, *, max_chars: int, max_bytes: int) -> tuple[str, bool]:
    """Plain-text/markdown body: a size-capped byte read decoded errors='replace'.

    Reads at most ``max_chars * 4`` bytes (UTF-8 worst case) so a huge log/txt is
    never slurped whole; a read that hits the cap marks the body truncated.
    """
    _guard_size(path, max_bytes)
    read_cap = max_chars * 4 + 8
    try:
        with open(path, "rb") as fh:
            raw = fh.read(read_cap + 1)
    except OSError as exc:
        raise DocumentError(f"cannot read text file: {exc}") from exc
    hit_cap = len(raw) > read_cap
    text = raw[:read_cap].decode("utf-8", errors="replace")
    return _normalize_body_text(text, max_chars, hard_stopped=hit_cap)


def extract_body(
    path: str,
    *,
    max_chars: int = DEFAULT_BODY_TEXT_MAX_CHARS,
    max_bytes: int,
    decompressed_max: int = DEFAULT_DECOMPRESSED_MAX,
    ratio_limit: float = DEFAULT_DECOMPRESSION_RATIO,
    ratio_min_bytes: int = DEFAULT_DECOMPRESSION_RATIO_MIN_BYTES,
) -> dict[str, Any]:
    """Extract plain body text for a supported document, or ``{}``.

    Dispatch by extension: pdf (pypdf), docx (python-docx, guarded), txt/md (byte
    read). xlsx/other spreadsheets get NO body text (structure only). Returns
    ``{"body_text": str, "body_text_truncated": bool}`` when non-empty text is
    produced, else ``{}``. Raises DocumentError on a hard parse failure or a
    decompression-guard rejection (the caller records ``_extract_error`` without
    discarding the already-parsed properties).
    """
    ext = PurePath(path).suffix.lstrip(".").lower()
    if ext == "pdf":
        body, truncated = _pdf_body(path, max_chars=max_chars, max_bytes=max_bytes)
    elif ext == "docx":
        # Guard BEFORE python-docx opens the archive (bomb rejected pre-parse).
        guard_decompression(
            path,
            decompressed_max=decompressed_max,
            ratio_limit=ratio_limit,
            ratio_min_bytes=ratio_min_bytes,
        )
        body, truncated = _docx_body(path, max_chars=max_chars, max_bytes=max_bytes)
    elif ext in ("txt", "md"):
        body, truncated = _text_body(path, max_chars=max_chars, max_bytes=max_bytes)
    else:
        return {}
    if not body:
        return {}
    return {"body_text": body, "body_text_truncated": truncated}


# Extension → parser. Anything not listed gets an ``unsupported`` marker (no
# error) because there is no safe metadata reader wired up for it in v1.
_PARSERS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "xlsx": extract_xlsx,
}
_ZIP_PARSERS = frozenset({"docx", "xlsx"})


def extract_document(
    path: str,
    *,
    max_bytes: int,
    decompressed_max: int = DEFAULT_DECOMPRESSED_MAX,
    ratio_limit: float = DEFAULT_DECOMPRESSION_RATIO,
    ratio_min_bytes: int = DEFAULT_DECOMPRESSION_RATIO_MIN_BYTES,
) -> dict[str, Any]:
    """Dispatch a document/spreadsheet path to the right property parser.

    Raises DocumentError on parse failure. Unknown/unsupported extensions return
    ``{"unsupported": True}`` rather than raising. Body text is a SEPARATE pass
    (:func:`extract_body`), merged by the ``extract.py`` wrapper.
    """
    ext = PurePath(path).suffix.lstrip(".").lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        return {"unsupported": True}
    if ext in _ZIP_PARSERS:
        return parser(
            path,
            max_bytes=max_bytes,
            decompressed_max=decompressed_max,
            ratio_limit=ratio_limit,
            ratio_min_bytes=ratio_min_bytes,
        )
    return parser(path, max_bytes=max_bytes)
